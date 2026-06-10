"""
Generate full documentation DOCX for Royal Glass Cost Calculator WordPress Plugin.
Run: python generate_docs.py
Output: RG_Calculator_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ──────────────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────────────

NAVY   = RGBColor(0x14, 0x2A, 0x4E)
BLUE   = RGBColor(0x1E, 0x5C, 0xB3)
AMBER  = RGBColor(0xB4, 0x56, 0x09)
GREEN  = RGBColor(0x18, 0x73, 0x4A)
GREY   = RGBColor(0x59, 0x59, 0x59)
LIGHT  = RGBColor(0xF0, 0xF4, 0xFA)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED    = RGBColor(0xC0, 0x00, 0x00)


def set_cell_bg(cell, hex_color: str):
    """Set background fill for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def h1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    run = p.runs[0]
    run.font.color.rgb = WHITE
    # Shade the paragraph via direct XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '142A4E')
    pPr.append(shd)
    return p


def h2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2')
    p.runs[0].font.color.rgb = NAVY
    return p


def h3(doc, text):
    p = doc.add_paragraph(text, style='Heading 3')
    p.runs[0].font.color.rgb = BLUE
    return p


def h4(doc, text):
    p = doc.add_paragraph(text, style='Heading 4')
    p.runs[0].font.color.rgb = GREY
    return p


def body(doc, text):
    return doc.add_paragraph(text, style='Normal')


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.add_run(text)
    return p


def numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.add_run(text)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4FA')
    pPr.append(shd)
    return p


def note_box(doc, text, color='F0F4FA', label='NOTE'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(f'  {label}:  {text}  ')
    run.font.size = Pt(10)
    run.font.italic = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    return p


def two_col_table(doc, rows, header=None, col_widths=None):
    cols = 2
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    table.style = 'Table Grid'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    offset = 0
    if header:
        for j, h in enumerate(header):
            cell = table.cell(0, j)
            set_cell_bg(cell, '142A4E')
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
        offset = 1
    for i, (key, val) in enumerate(rows):
        r = i + offset
        c0 = table.cell(r, 0)
        c1 = table.cell(r, 1)
        if i % 2 == 0:
            set_cell_bg(c0, 'F0F4FA')
            set_cell_bg(c1, 'F0F4FA')
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(key)
        run0.font.size = Pt(10)
        run0.bold = True
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(val)
        run1.font.size = Pt(10)
    return table


def page_break(doc):
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # ── Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Style defaults
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for hn, sz, bold in [('Heading 1', 16, True), ('Heading 2', 14, True),
                          ('Heading 3', 12, True), ('Heading 4', 11, True)]:
        s = doc.styles[hn]
        s.font.name = 'Calibri'
        s.font.size = Pt(sz)
        s.font.bold = bold
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after  = Pt(6)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # COVER PAGE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('Royal Glass Cost Calculator')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('WordPress Plugin — Full Documentation')
    run.font.size = Pt(16)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_p.add_run(f'Version 2.3.0   ·   Generated {datetime.date.today().strftime("%d %B %Y")}')
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    run.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    note_box(doc,
        'This document covers both non-technical users (site owners, admin staff) and developers. '
        'If you are a site owner, focus on Parts 1–3. If you are a developer, all parts apply.',
        color='E8F4FD', label='WHO THIS IS FOR')

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # TABLE OF CONTENTS (manual)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'Table of Contents')
    toc_entries = [
        ('PART 1', 'Overview & What This Plugin Does'),
        ('PART 2', 'Using the Calculator (Customer Guide)'),
        ('PART 3', 'WordPress Admin Guide'),
        ('  3.1', 'Managing Leads'),
        ('  3.2', 'Editing Pricing'),
        ('PART 4', 'Installation & Setup'),
        ('  4.1', 'Requirements'),
        ('  4.2', 'Fresh Installation'),
        ('  4.3', 'Updating the Plugin'),
        ('  4.4', 'Configuration (wp-config.php)'),
        ('PART 5', 'Developer Reference'),
        ('  5.1', 'Repository Structure'),
        ('  5.2', 'React App Architecture'),
        ('  5.3', 'The 9-Step Wizard Flow'),
        ('  5.4', 'Pricing Engine'),
        ('  5.5', 'REST API Endpoints'),
        ('  5.6', 'Database Schema'),
        ('  5.7', 'Email System'),
        ('  5.8', 'ServiceM8 Integration'),
        ('  5.9', 'Security Model'),
        ('  5.10', 'Build & Deploy Commands'),
        ('PART 6', 'Common Issues & Troubleshooting'),
        ('PART 7', 'Known Limitations & Future Work'),
        ('PART 8', 'Quick Reference'),
    ]
    for num, title in toc_entries:
        p = doc.add_paragraph()
        tab_run = p.add_run(f'{num}')
        tab_run.font.bold = True if not num.startswith('  ') else False
        tab_run.font.size = Pt(11)
        p.add_run(f'   {title}').font.size = Pt(11)

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 1 — OVERVIEW
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 1 — Overview & What This Plugin Does')

    body(doc,
        'The Royal Glass Cost Calculator is a WordPress plugin that gives website visitors an '
        'instant price estimate for frameless glass balustrades and pool fences. '
        'The visitor answers a short series of questions, provides their contact details, and '
        'receives an estimate on screen and by email — without waiting for a call-back.')

    doc.add_paragraph()
    h2(doc, 'What happens when someone uses it')
    numbered(doc, 'Visitor goes to royalglass.co.nz/estimate.')
    numbered(doc, 'They answer up to 9 questions about their project (type, length, materials, etc.).')
    numbered(doc, 'They enter their contact details and consent to being contacted.')
    numbered(doc, 'They see their price range instantly on screen.')
    numbered(doc, 'Three things happen automatically in the background:')
    bullet(doc, 'Admin receives a plain-text email with all project and contact details.', level=1)
    bullet(doc, 'ServiceM8 inbox receives an email ready to "Convert to Job".', level=1)
    bullet(doc, 'Customer receives a branded HTML email with their estimate.', level=1)
    numbered(doc, 'The lead appears in WordPress Admin → RG Calculator → Leads for follow-up.')

    doc.add_paragraph()
    h2(doc, 'Technology summary')
    two_col_table(doc, [
        ('Frontend',    'React (TypeScript) single-page app, built with Vite'),
        ('Hosting',     'WordPress plugin shortcode — no separate server needed'),
        ('Database',    'WordPress MySQL — custom table wp_rg_leads'),
        ('Emails',      'WordPress wp_mail() (uses server mail config or SMTP plugin)'),
        ('Bot protection', 'Cloudflare Turnstile CAPTCHA + honeypot + time gate'),
        ('Address lookup', 'OpenStreetMap Nominatim (free, no API key required)'),
        ('CRM link',    'ServiceM8 via inbox email parsing'),
    ], header=['Component', 'Technology'])

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 2 — CUSTOMER GUIDE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 2 — Using the Calculator (Customer Guide)')

    body(doc,
        'This section explains what customers see and do at each step. Share this with front-line '
        'staff so they can help customers who call with questions.')

    h2(doc, 'Step-by-step walkthrough')

    steps = [
        ('Step 1 — What is your project?',
         'Choose the type of glass installation:\n'
         '  • Ground Level Balustrade — flat area at ground level\n'
         '  • Balcony / Patio Balustrade — elevated area (deck, balcony)\n'
         '  • Premium Pool Fence — meets NZ Pool Safety Act requirements\n'
         '  • Stair Balustrade — along a staircase\n\n'
         'Tip: If unsure, choose the closest match — the estimate will still be valid for quoting.'),
        ('Step 2 — Total length of glass run',
         'Move the slider to enter the total length in metres.\n'
         'Range: 1 m to 100 m. Jobs under 5 m are charged at the 5 m minimum.\n'
         'For stair balustrades: enter the stair run length and any landing length separately.'),
        ('Step 3 — How many corners? (not shown for stairs)',
         'Enter the number of 90° turns. A straight run has 0 corners. '
         'Each corner adds a small surcharge to account for additional brackets and cuts.'),
        ('Step 4 — How many gates? (pool fence only)',
         'Enter the number of gates needed. Pool Safety rules require at least one self-closing, '
         'self-latching gate — a warning appears if 0 is selected. Each gate has a fixed price.'),
        ('Step 5 — Glass type (balcony & stairs only)',
         '  • 12 mm Toughened + Capping — the standard, most common choice\n'
         '  • Laminated — higher impact resistance, required in some commercial settings\n\n'
         'Pool fences always use 12 mm toughened. Ground level uses the standard option automatically.'),
        ('Step 6 — Glass colour',
         '  • Clear — standard, most popular\n'
         '  • Tinted — reduces glare, adds privacy\n'
         '  • Frosted — maximum privacy\n'
         '  • Low-Iron (Starphire) — ultra-clear, slightly premium price'),
        ('Step 7 — How will the glass be fixed?',
         '  • Spigot Round / Standoff Posts — floor-mounted posts\n'
         '  • Viking — surface-mounted frameless system\n'
         '  • JH Clamps — side-fix clamps\n'
         '  • Side Channel / Top Channel — channel systems\n'
         '  • Aluminium (Type 1 or 2) — aluminium frame systems\n'
         '  • SED (Special Engineer Design) — non-standard, requires engineer sign-off\n\n'
         'Not sure? Select "Not sure" — a note will appear on the estimate.'),
        ('Step 8 — What is the substrate?',
         'What is the surface the glass will be fixed to?\n'
         '  • Timber — deck boards or timber framing\n'
         '  • Concrete — slab or concrete structure\n'
         '  • Tile — tiled surface (may need special fixings)\n'
         '  • Steel — steel framing or posts\n'
         '  • Not sure — a note will appear on the estimate for site confirmation'),
        ('Step 9 — Hardware finish',
         '  • Standard Chrome — included in base price\n'
         '  • Matte Black — popular contemporary finish, small surcharge\n'
         '  • Brushed Chrome — subtle metallic finish, small surcharge\n'
         '  • Powder Coated — custom colour matching, slightly higher surcharge\n'
         '  • Not sure — a note will appear on the estimate'),
        ('Contact form',
         'After the calculator steps, the customer fills in:\n'
         '  • Full name (or company name)\n'
         '  • Email address\n'
         '  • Phone number (NZ format)\n'
         '  • Customer type (homeowner, builder, developer, architect, pool builder, other)\n'
         '  • Timeframe (ASAP / 1–3 months / 3–6 months / 6+ months / just planning)\n'
         '  • Project address (autocompletes from OpenStreetMap)\n'
         '  • Optional notes\n'
         '  • Consent checkbox (required)\n\n'
         'They then click "Get my estimate" to submit.'),
        ('Result screen',
         'Shows:\n'
         '  • Price range (low–high, excl. GST)\n'
         '  • Amber note if any details need site confirmation\n'
         '  • Full project summary table\n'
         '  • "Share this estimate" — lets them forward to a builder or partner\n'
         '  • Call-to-action: phone number and what to expect next'),
    ]
    for title, desc in steps:
        h3(doc, title)
        for line in desc.split('\n'):
            if line.startswith('  •'):
                bullet(doc, line[3:], level=0)
            elif line.strip():
                body(doc, line)

    h2(doc, 'What the customer receives by email')
    body(doc,
        'Immediately after submitting, the customer receives a branded HTML email containing:')
    bullet(doc, 'Their price range (large, prominent)')
    bullet(doc, 'A note about what the estimate includes and excludes')
    bullet(doc, 'A 3-step timeline: "We\'ll call within 1 business day → Site visit → Written quote"')
    bullet(doc, 'Their full project summary')
    bullet(doc, 'Any items flagged for site confirmation (amber boxes)')
    bullet(doc, 'Royal Glass contact details')
    note_box(doc,
        'The estimate email is sent automatically. Customers should check spam if they don\'t see it within a few minutes.',
        color='FFF8E1', label='TIP')

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 3 — ADMIN GUIDE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 3 — WordPress Admin Guide')

    body(doc,
        'Log in to your WordPress dashboard. In the left-hand menu you will see '
        '"RG Calculator" with two sub-pages: Leads and Pricing.')

    h2(doc, '3.1 — Managing Leads')

    h3(doc, 'The Leads list page')
    body(doc, 'Shows all submitted enquiries. Use the status tabs at the top to filter:')
    two_col_table(doc, [
        ('NEW',      'Just received — not yet acted on'),
        ('REVIEWED', 'You have opened and read the lead'),
        ('ACCEPTED', 'You have agreed to quote / follow up'),
        ('REJECTED', 'Spam or out-of-scope — archived'),
    ], header=['Status', 'Meaning'], col_widths=[1.5, 4.5])

    doc.add_paragraph()
    body(doc, 'Each row shows: lead number, name, phone, email, project type, estimate range, status, and date.')
    body(doc, 'Click a lead row to open the detail view.')

    h3(doc, 'Lead detail view')
    body(doc, 'Shows all information the customer provided:')
    bullet(doc, 'Contact details: name, phone, email, address, customer type, preferred call time')
    bullet(doc, 'Project details: scenario, length, corners, gates, glass type & colour, fixing method, substrate, hardware')
    bullet(doc, 'Estimate: low/high range')
    bullet(doc, 'Any site consultation flags (amber notes)')
    bullet(doc, 'Customer notes (optional, free text)')
    bullet(doc, 'ServiceM8 sync status and timestamp')
    body(doc, 'Use the status buttons (REVIEWED / ACCEPTED / REJECTED) to update the lead status.')

    h3(doc, 'Lead workflow recommendation')
    numbered(doc, 'New lead arrives — admin email received (check inbox).')
    numbered(doc, 'Open in WP Admin → mark REVIEWED.')
    numbered(doc, 'Call or email the customer within 1 business day.')
    numbered(doc, 'Mark ACCEPTED (proceeding with quote) or REJECTED (not suitable).')
    numbered(doc, 'In ServiceM8: find the email in the inbox, click "Convert to Job" to create a full job record.')

    note_box(doc,
        'Leads are never deleted automatically. Archive rejected leads so the list stays manageable.',
        color='FFF8E1', label='TIP')

    doc.add_paragraph()
    h2(doc, '3.2 — Editing Pricing')

    body(doc, 'Go to RG Calculator → Pricing to change any price without touching code.')

    h3(doc, 'Pricing fields explained')
    two_col_table(doc, [
        ('Base rate per metre',         'The $/m rate for each project type. Main lever for profitability.'),
        ('Gate price',                  'Fixed price per gate (pool fence only).'),
        ('Minimum length',              'Jobs shorter than this are charged as if this long (default: 5 m).'),
        ('Corner surcharge',            'Added per 90° corner (default: $85).'),
        ('Hardware finish surcharges',  'Per-metre add-on for non-chrome finishes (matte black, brushed, powder).'),
        ('Glass type surcharges',       'Per-metre add-on for laminated glass over standard toughened.'),
        ('Glass colour surcharges',     'Per-metre add-on for tinted, frosted, or low-iron glass.'),
        ('Fixing method surcharges',    'Per-metre add-on (or discount with a negative value) for each fixing type.'),
        ('Interlinking rails surcharge','Per-metre surcharge when rails are included.'),
        ('Price range low/high %',      'The estimate shows a band (e.g. 90%–120% of subtotal). Adjust to widen or narrow the range.'),
    ], header=['Field', 'What it does'], col_widths=[2.5, 3.5])

    doc.add_paragraph()
    note_box(doc,
        'Changes take effect immediately — no rebuild required. The React app loads pricing from '
        'the WordPress REST API on page load.',
        color='E8F8E8', label='IMPORTANT')

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 4 — INSTALLATION & SETUP
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 4 — Installation & Setup')

    h2(doc, '4.1 — Requirements')
    two_col_table(doc, [
        ('WordPress',   '5.8 or later (tested on 6.x)'),
        ('PHP',         '7.4 or later (8.x recommended)'),
        ('MySQL',       '5.7 or later'),
        ('Hosting',     'Bluehost shared hosting (or equivalent). No Cloudflare proxy in front.'),
        ('Email',       'Server PHP mail() or SMTP plugin (e.g. WP Mail SMTP). Gmail SMTP recommended for reliability.'),
        ('Browser',     'Chrome, Firefox, Edge, Safari — last 2 major versions.'),
    ], header=['Requirement', 'Details'], col_widths=[2.0, 4.0])

    doc.add_paragraph()
    h2(doc, '4.2 — Fresh Installation')
    numbered(doc, 'Build the plugin ZIP (developer step — see Part 5.10), or obtain a pre-built ZIP.')
    numbered(doc, 'WordPress Admin → Plugins → Add New → Upload Plugin.')
    numbered(doc, 'Select the ZIP file → Install Now → Activate Plugin.')
    numbered(doc, 'Create the calculator page:')
    bullet(doc, 'Pages → Add New', level=1)
    bullet(doc, 'Title: "Get an Instant Estimate" (or any name)', level=1)
    bullet(doc, 'Slug: estimate (so the URL is royalglass.co.nz/estimate)', level=1)
    bullet(doc, 'Content: add the shortcode  [rg_calculator]', level=1)
    bullet(doc, 'Page template: Full Width (hides the sidebar for a clean layout)', level=1)
    bullet(doc, 'Publish', level=1)
    numbered(doc, 'Fix permalinks if the calculator doesn\'t load:')
    bullet(doc, 'Settings → Permalinks → select "Post name" → Save Changes.', level=1)
    numbered(doc, 'Add the configuration constants to wp-config.php (see section 4.4).')
    numbered(doc, 'Verify: visit the estimate page — the calculator should appear.')

    h2(doc, '4.3 — Updating the Plugin')

    h3(doc, 'JavaScript/CSS only (no PHP changes)')
    body(doc, 'This is the most common update. No plugin deactivation needed.')
    numbered(doc, 'Build the React app (npm run build).')
    numbered(doc, 'Upload the two new files via SFTP or cPanel File Manager:')
    bullet(doc, 'dist/rg-calculator.js  →  wp-content/plugins/rg-calculator/assets/', level=1)
    bullet(doc, 'dist/rg-calculator.css →  wp-content/plugins/rg-calculator/assets/', level=1)
    numbered(doc, 'Hard-refresh the calculator page to verify (Ctrl+Shift+R on Windows).')
    note_box(doc,
        'Do NOT upload any .jpg files from dist/. Images live permanently in the assets folder and must not be overwritten.',
        color='FFE8E8', label='WARNING')

    h3(doc, 'PHP changes included')
    numbered(doc, 'Build the full plugin ZIP.')
    numbered(doc, 'Deactivate the existing plugin (WordPress Admin → Plugins → Deactivate).')
    numbered(doc, 'Delete the old plugin files (keep a backup first).')
    numbered(doc, 'Upload and activate the new ZIP.')

    h2(doc, '4.4 — Configuration (wp-config.php)')
    body(doc, 'Add these constants to wp-config.php (before the "That\'s all, stop editing" comment):')
    code_block(doc,
        "// Cloudflare Turnstile bot protection\n"
        "// Get keys from: dash.cloudflare.com → Turnstile → Add Site\n"
        "define('RG_TURNSTILE_SITE_KEY', '0x4AAA...');  // public key\n"
        "define('RG_TURNSTILE_SECRET',   '0x4AAA...');  // private key\n\n"
        "// Who receives admin lead notification emails\n"
        "define('RG_LEAD_NOTIFY_EMAIL', 'sales@royalglass.co.nz');\n\n"
        "// ServiceM8 inbox email (leave undefined to disable SM8 emails)\n"
        "define('RG_SM8_INBOX_EMAIL', 'xxxxxx@inbox.servicem8.com');\n\n"
        "// Google Maps key (currently unused — address is Nominatim-based)\n"
        "// define('RG_GOOGLE_MAPS_KEY', 'AIza...');"
    )

    two_col_table(doc, [
        ('RG_TURNSTILE_SITE_KEY',  'Public Cloudflare key — goes into the React frontend. Safe to expose.'),
        ('RG_TURNSTILE_SECRET',    'Private Cloudflare key — backend only. Keep secret.'),
        ('RG_LEAD_NOTIFY_EMAIL',   'Email that receives admin lead notifications. Falls back to WP admin email if not set.'),
        ('RG_SM8_INBOX_EMAIL',     'ServiceM8 inbox address. Leave undefined to disable. Accepts comma-separated list.'),
    ], header=['Constant', 'Description'], col_widths=[2.5, 3.5])

    note_box(doc,
        'Both Turnstile keys must be defined for bot-checking to activate. '
        'Defining only one treats the CAPTCHA as disabled.',
        color='FFF8E1', label='NOTE')

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 5 — DEVELOPER REFERENCE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 5 — Developer Reference')

    # ── 5.1 REPOSITORY STRUCTURE
    h2(doc, '5.1 — Repository Structure')
    code_block(doc,
        "cost-calculator-WP/\n"
        "├── src/                          React + TypeScript source\n"
        "│   ├── main.tsx                  Entry — mounts on #rg-calculator-root\n"
        "│   ├── App.tsx                   State root — wizard answers + screen routing\n"
        "│   ├── index.css                 Scoped reset under #rg-calculator-root\n"
        "│   ├── components/\n"
        "│   │   ├── CalculatorForm.tsx    Steps 1–7 wizard + navigation logic\n"
        "│   │   └── wizard/\n"
        "│   │       ├── LeadCapture.tsx   Step 8 — contact form, Turnstile, submission\n"
        "│   │       ├── ResultScreen.tsx  Step 9 — estimate display + share panel\n"
        "│   │       ├── NZAddressAutocomplete.tsx  Nominatim address lookup\n"
        "│   │       └── steps/shared.tsx  SelectionCard, SliderInput, StepNote, etc.\n"
        "│   ├── hooks/\n"
        "│   │   └── usePricing.ts         Fetches live pricing from REST API on mount\n"
        "│   └── lib/calculator/\n"
        "│       ├── types.ts              WizardAnswers, LeadData, EstimateResult, PricingConfig\n"
        "│       ├── config.ts             DEFAULT_PRICING, IMAGES map, getPluginBase()\n"
        "│       └── engine.ts             calculateEstimate() — pure, no side effects\n"
        "├── wordpress-plugin/rg-calculator/\n"
        "│   ├── rg-calculator.php         Plugin bootstrap, shortcode, SEO schema\n"
        "│   ├── assets/\n"
        "│   │   ├── rg-calculator.js      Built React bundle (copy from dist/)\n"
        "│   │   ├── rg-calculator.css     Built styles (copy from dist/)\n"
        "│   │   └── *.jpg                 Wizard option images — never overwrite\n"
        "│   └── includes/\n"
        "│       ├── database.php          wp_rg_leads schema + CRUD helpers\n"
        "│       ├── validation.php        Input validation & sanitization\n"
        "│       ├── api.php               REST routes: /leads, /pricing, /estimate-email\n"
        "│       ├── email.php             Admin + customer HTML emails\n"
        "│       ├── servicem8.php         SM8 inbox email + cron retry\n"
        "│       ├── admin-pricing.php     WP admin pricing editor\n"
        "│       └── admin-leads.php       WP admin leads list & detail\n"
        "├── package.json\n"
        "├── vite.config.ts\n"
        "├── CLAUDE.md                     Architectural decisions & constraints\n"
        "└── CHANGELOG.md"
    )

    # ── 5.2 REACT APP ARCHITECTURE
    h2(doc, '5.2 — React App Architecture')
    body(doc,
        'The app is a single-page React/TypeScript app built with Vite. It is delivered as two static '
        'files (JS bundle + CSS) loaded by the WordPress plugin.')

    h3(doc, 'Styling rule — NO Tailwind, inline styles only')
    body(doc,
        'All component styling uses inline style props. Tailwind has been permanently removed. '
        'WordPress themes output unlayered CSS that always overrides Tailwind\'s @layer utilities. '
        'Inline styles are immune to the cascade. Never reintroduce Tailwind className props.')

    h3(doc, 'State management')
    body(doc,
        'All wizard state lives in App.tsx local React state. No Redux, Zustand, or Context API. '
        'Pricing is fetched once on mount via usePricing.ts and passed down as props.')

    h3(doc, 'Image assets')
    body(doc,
        'Images are NOT imported as ES modules. They are referenced via the IMAGES map in config.ts, '
        'which resolves URLs from window.rgCalculatorConfig.assetsUrl (injected by WordPress via '
        'wp_localize_script). Never add image imports or rely on Vite\'s base URL for images.')

    h3(doc, 'window.rgCalculatorConfig')
    body(doc, 'WordPress injects this global object before the React bundle loads:')
    code_block(doc,
        "window.rgCalculatorConfig = {\n"
        "  restUrl:         'https://royalglass.co.nz/wp-json/royal-glass/v1/', // REST base\n"
        "  nonce:           'abc123...',    // wp_create_nonce('wp_rest')\n"
        "  googleMapsKey:   '',             // legacy, currently unused\n"
        "  turnstileSiteKey:'0x4AAA...',    // Cloudflare site key (or '')\n"
        "  assetsUrl:       'https://royalglass.co.nz/wp-content/plugins/rg-calculator/assets/'\n"
        "};"
    )

    # ── 5.3 WIZARD FLOW
    h2(doc, '5.3 — The 9-Step Wizard Flow')
    body(doc, 'Steps 1–7 live in CalculatorForm.tsx. Steps 8–9 are LeadCapture.tsx and ResultScreen.tsx.')

    wizard_rows = [
        ('1', 'Scenario (project type)', 'Always shown', 'ground_level | balcony_balustrade | premium_pool_fence | stair_balustrade'),
        ('2', 'Total length (metres)', 'Always shown', 'Slider 1–100 m. Stairs split into run + landing.'),
        ('3', 'Corners', 'Hidden for stair_balustrade', 'Slider 0–10'),
        ('4', 'Gates', 'Only for premium_pool_fence', 'Slider 0–6. NZ Pool Safety warning if 0.'),
        ('5', 'Glass type', 'Hidden for ground_level & premium_pool_fence', 'toughened_12mm | laminated'),
        ('6', 'Glass colour', 'Always shown', 'clear | tinted | frosted | low_iron'),
        ('7', 'Fixing method', 'Always shown', 'spigot_round | standoff_posts | viking | jh_clamps | side_channel | top_channel | aluminium_1 | aluminium_2 | sed'),
        ('8', 'Substrate', 'Always shown — mandatory before Continue', 'timber | concrete | tile | steel | not_sure'),
        ('9', 'Hardware finish', 'Always shown', 'standard_chrome | matte_black | brushed_chrome | powder_coated | not_sure'),
    ]
    table = doc.add_table(rows=len(wizard_rows) + 1, cols=4)
    table.style = 'Table Grid'
    for j, h in enumerate(['Step', 'Name', 'Condition', 'Valid values']):
        cell = table.cell(0, j)
        set_cell_bg(cell, '142A4E')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
    for i, row in enumerate(wizard_rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            if i % 2 == 0:
                set_cell_bg(cell, 'F0F4FA')
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()
    h3(doc, 'Consultation flags (never block the estimate — informational only)')
    two_col_table(doc, [
        ('fixing_method = not_sure',   'Fixing method to be confirmed on site'),
        ('fixing_method = sed',        'Special Engineer Design required'),
        ('hardware = not_sure',        'Hardware finish to be confirmed'),
        ('substrate = not_sure',       'Substrate to be confirmed on site'),
    ], header=['Condition', 'Flag text shown to customer'], col_widths=[2.5, 3.5])

    # ── 5.4 PRICING ENGINE
    h2(doc, '5.4 — Pricing Engine')

    h3(doc, 'Where prices are defined')
    body(doc,
        'src/lib/calculator/config.ts contains DEFAULT_PRICING — the fallback used when the '
        'REST API call fails or during local development. Live pricing is stored in the '
        'rg_calculator_pricing WordPress option and served by GET /pricing.')

    h3(doc, 'Calculation formula (engine.ts: calculateEstimate())')
    code_block(doc,
        "effectiveLength  = max(inputLength, minimumLength)\n"
        "base             = effectiveLength × scenario.ratePerMetre\n"
        "gates            = gateCount × scenario.gatePrice    (pool fence only)\n"
        "corners          = cornerCount × cornerSurcharge\n"
        "hardware         = effectiveLength × hardwareFinishSurcharge[hardwareFinish]\n"
        "glassType        = effectiveLength × glassTypeSurcharge[glassType]\n"
        "glassColour      = effectiveLength × glassColourSurcharge[glassColour]\n"
        "rails            = interlikingRails ? effectiveLength × interlikingRailsSurcharge : 0\n"
        "fixingMethod     = effectiveLength × fixingMethodSurcharge[fixingMethod]\n\n"
        "subtotal         = sum of all above\n"
        "low              = round(subtotal × (rangeLowPercent  / 100), to nearest $50)\n"
        "high             = round(subtotal × (rangeHighPercent / 100), to nearest $50)"
    )

    h3(doc, 'Default pricing values')
    two_col_table(doc, [
        ('Ground Level — base rate',    '$280 / m'),
        ('Balcony Balustrade',          '$320 / m'),
        ('Premium Pool Fence',          '$380 / m + $680 / gate'),
        ('Stair Balustrade',            '$330 / m'),
        ('Minimum length',              '5 m'),
        ('Corner surcharge',            '$85 / corner'),
        ('Matte Black hardware',        '+$15 / m'),
        ('Brushed Chrome hardware',     '+$12 / m'),
        ('Powder Coated hardware',      '+$22 / m'),
        ('All glass type surcharges',   '$0 / m (admin-configurable)'),
        ('All glass colour surcharges', '$0 / m (admin-configurable)'),
        ('All fixing method surcharges','$0 / m (admin-configurable, negatives allowed)'),
        ('Price range',                 '90% – 120% of subtotal'),
    ], header=['Item', 'Default'], col_widths=[3.0, 3.0])

    # ── 5.5 REST API
    h2(doc, '5.5 — REST API Endpoints')
    body(doc, 'Base URL: /wp-json/royal-glass/v1')

    h3(doc, 'GET /pricing')
    two_col_table(doc, [
        ('Permission',  'Public (no auth required)'),
        ('Response',    '200 JSON — full PricingConfig object'),
        ('Use',         'Called by React app on mount. Falls back to DEFAULT_PRICING on failure.'),
    ], col_widths=[2.0, 4.0])

    h3(doc, 'POST /leads')
    two_col_table(doc, [
        ('Permission',  'Public'),
        ('Body',        'JSON: { answers, lead, estimate, turnstileToken, loadedAt }'),
        ('Security checks (in order)',
         '1) Time gate (< 3 s → 429)\n'
         '2) Honeypot (websiteUrl filled → 400)\n'
         '3) Turnstile verify (if both keys set)\n'
         '4) Rate limit (10/IP/hour)'),
        ('Validation',  'rg_validate_lead() + rg_validate_answers() + rg_sanitize_estimate()'),
        ('Success',     '201  { "ok": true, "leadId": 42 }'),
        ('Side effects','Saves to DB, then on shutdown: admin email + SM8 email + customer email'),
    ], col_widths=[2.0, 4.0])

    h3(doc, 'POST /estimate-email')
    two_col_table(doc, [
        ('Permission',  'Public — requires WP nonce header X-WP-Nonce'),
        ('Body',        'JSON: { email, firstName, leadId, answers, estimate }'),
        ('Security',    'wp_verify_nonce + rate limit (5/IP/hour)'),
        ('Validation',  'rg_validate_answers() called before sending'),
        ('Success',     '200  { "ok": true }'),
        ('Side effects','Sends customer estimate email to provided address asynchronously'),
    ], col_widths=[2.0, 4.0])

    h3(doc, 'Error response codes')
    two_col_table(doc, [
        ('400', 'Invalid JSON, honeypot triggered, or bad request'),
        ('403', 'Turnstile verification failed, or invalid nonce'),
        ('422', 'Validation failure — response body contains { "error": "..." }'),
        ('429', 'Rate limit exceeded or form submitted too quickly'),
        ('500', 'Database save failed'),
    ], header=['Code', 'Meaning'], col_widths=[1.0, 5.0])

    # ── 5.6 DATABASE SCHEMA
    h2(doc, '5.6 — Database Schema')
    body(doc, 'Table: wp_rg_leads (prefix may vary). Created on plugin activation.')

    db_rows = [
        ('id',                  'BIGINT AUTO_INCREMENT',     'Primary key'),
        ('status',              'VARCHAR(20) DEFAULT "NEW"', 'NEW | REVIEWED | ACCEPTED | REJECTED'),
        ('first_name',          'VARCHAR(100)',              ''),
        ('last_name',           'VARCHAR(100)',              ''),
        ('phone',               'VARCHAR(50)',               'NZ format validated'),
        ('email',               'VARCHAR(255)',              ''),
        ('address',             'TEXT',                      ''),
        ('customer_type',       'VARCHAR(30)',               'homeowner | builder | developer | architect | pool_builder | other'),
        ('call_pref',           'VARCHAR(20)',               'Default: anytime'),
        ('notes',               'TEXT',                      'Customer-supplied optional notes'),
        ('project_type',        'VARCHAR(50)',               'Wizard scenario value'),
        ('length_m',            'SMALLINT',                  'Effective metres (post-minimum)'),
        ('corners',             'TINYINT',                   ''),
        ('gates',               'TINYINT',                   ''),
        ('fixing_method',       'VARCHAR(50)',               ''),
        ('substrate',           'VARCHAR(50)',               ''),
        ('hardware',            'VARCHAR(50)',               ''),
        ('est_low',             'DECIMAL(10,2)',             'NZD, excl. GST'),
        ('est_high',            'DECIMAL(10,2)',             'NZD, excl. GST'),
        ('est_subtotal',        'DECIMAL(10,2)',             'Pre-band subtotal'),
        ('needs_consult',       'TINYINT(1)',                'Boolean — any consultation flags'),
        ('consult_notes',       'TEXT',                      'Serialized flag list'),
        ('consent_given',       'TINYINT(1)',                'Required contact consent'),
        ('consented_at',        'DATETIME',                  'Timestamp of consent'),
        ('servicem8_status',    'VARCHAR(20)',               'sent_to_inbox | failed | skipped'),
        ('servicem8_sent_at',   'DATETIME',                  ''),
        ('created_at',          'DATETIME',                  'Submission timestamp'),
        ('updated_at',          'DATETIME',                  'Auto-updates on status change'),
    ]
    table = doc.add_table(rows=len(db_rows) + 1, cols=3)
    table.style = 'Table Grid'
    for j, h in enumerate(['Column', 'Type', 'Notes']):
        cell = table.cell(0, j)
        set_cell_bg(cell, '142A4E')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
    for i, row in enumerate(db_rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            if i % 2 == 0:
                set_cell_bg(cell, 'F0F4FA')
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if j == 0:
                run.font.name = 'Courier New'

    # ── 5.7 EMAIL SYSTEM
    doc.add_paragraph()
    h2(doc, '5.7 — Email System')

    body(doc, 'Three emails are sent asynchronously on every lead submission:')

    h3(doc, 'How async sending works')
    body(doc,
        'All emails are queued via add_action("shutdown", ...) and fastcgi_finish_request() is called '
        'first. This means the browser receives the 201 response immediately; emails send in the '
        'background. ignore_user_abort(true) ensures sending completes even if the user navigates away.')

    h3(doc, 'Email 1 — Admin lead notification')
    two_col_table(doc, [
        ('To',      'RG_LEAD_NOTIFY_EMAIL constant, or WordPress admin email'),
        ('Format',  'Plain text'),
        ('Subject', '"New Enquiry Lead#{id} - {project} - {name}"'),
        ('Content', 'Contact details, project details, estimate range, notes, admin link'),
        ('Function','rg_send_lead_email()  in includes/email.php'),
    ], col_widths=[1.5, 4.5])

    h3(doc, 'Email 2 — ServiceM8 inbox')
    two_col_table(doc, [
        ('To',      'RG_SM8_INBOX_EMAIL constant. Disabled if constant not defined.'),
        ('Format',  'Plain text, structured for SM8 "Convert to Job" parser'),
        ('Subject', '"New Enquiry — {name} — {project} — Est. ${low}–${high}"'),
        ('Content', 'Name, Mobile, Email, Address fields SM8 can parse; project details; WP admin link'),
        ('Function','rg_sm8_send_immediate()  in includes/servicem8.php'),
    ], col_widths=[1.5, 4.5])

    h3(doc, 'Email 3 — Customer estimate email')
    two_col_table(doc, [
        ('To',      'Customer email address from lead form'),
        ('From',    '"Royal Glass Limited <support@royalglass.co.nz>"'),
        ('Format',  'Rich HTML'),
        ('Subject', '"{firstName}, your Royal Glass estimate is here"'),
        ('Content', 'Price range banner, disclaimer, 3-step timeline, CTA phone, project summary, consultation flags, footer'),
        ('Function','rg_send_estimate_email_to_customer()  in includes/email.php'),
    ], col_widths=[1.5, 4.5])

    note_box(doc,
        'Email deliverability depends on your server mail configuration. '
        'Install WP Mail SMTP and configure a Gmail/Sendgrid/Mailgun sender for reliable delivery.',
        color='FFF8E1', label='IMPORTANT')

    # ── 5.8 SERVICEM8
    h2(doc, '5.8 — ServiceM8 Integration')
    body(doc,
        'The integration is email-only — no direct API calls are made to ServiceM8. '
        'This avoids exposing SM8 credentials in a publicly accessible PHP file.')

    h3(doc, 'Flow')
    numbered(doc, 'Lead submitted → saved to DB.')
    numbered(doc, 'On shutdown hook, rg_sm8_send_immediate() fires.')
    numbered(doc, 'Plain-text email sent to RG_SM8_INBOX_EMAIL.')
    numbered(doc, 'Database columns servicem8_status + servicem8_sent_at are updated.')
    numbered(doc, 'In ServiceM8: staff sees email in inbox → clicks "Convert to Job" → SM8 auto-fills client details from the structured email.')

    h3(doc, 'Enabling / Disabling')
    bullet(doc, 'Define RG_SM8_INBOX_EMAIL in wp-config.php to enable.')
    bullet(doc, 'Remove or comment out the constant to disable — no emails will be sent.')
    bullet(doc, 'Accepts a comma-separated list for multiple recipients.')

    h3(doc, 'Cron retry')
    body(doc,
        'A WordPress cron job (every 15 minutes) retries leads marked pending_inbox older than 10 minutes, '
        'up to 3 attempts, before marking them failed. This handles transient mail failures.')

    # ── 5.9 SECURITY
    h2(doc, '5.9 — Security Model')

    h3(doc, 'Layers of protection on POST /leads')
    two_col_table(doc, [
        ('Time gate',          'Rejects submissions under 3 seconds — stops instant-submit bots'),
        ('Honeypot field',     'Hidden websiteUrl field — bots fill it in, humans don\'t'),
        ('Turnstile CAPTCHA',  'Cloudflare invisible CAPTCHA — verifies human interaction'),
        ('Rate limiting',      '10 submissions per IP per hour (MySQL transient-backed)'),
        ('Server validation',  'rg_validate_lead() + rg_validate_answers() — runs independently of frontend validation'),
        ('SQL injection',      '$wpdb->insert() with format specifiers — no raw interpolation'),
        ('Admin auth',         'All WP admin pages gated by current_user_can("manage_options") + nonce'),
    ], header=['Layer', 'What it does'], col_widths=[2.0, 4.0])

    h3(doc, 'Known limitations (low risk)')
    two_col_table(doc, [
        ('/estimate-email',    'Missing Turnstile + honeypot. Rate limit + nonce only. Low practical risk on Bluehost.'),
        ('Rate limit atomicity','Non-atomic check/set — concurrent requests can exceed limit by 1.'),
        ('Estimate values',    'Stored as-supplied from client (bounds-clamped but not server-recalculated).'),
        ('marketingConsent',   'Captured in React state but not sent to server or stored.'),
    ], header=['Item', 'Details'], col_widths=[2.0, 4.0])

    h3(doc, 'Non-negotiable constraints (never violate)')
    bullet(doc, 'Never reintroduce Tailwind className props — use inline style props only.')
    bullet(doc, 'Never add endpoints that skip rg_validate_answers() when answers payload is present.')
    bullet(doc, 'Never skip Turnstile verification on new public POST endpoints when both keys are configured.')
    bullet(doc, 'Never expose SM8 API keys, n8n webhooks, or Supabase keys in the browser bundle.')
    bullet(doc, 'Database save must complete before emails are scheduled — if DB insert fails, return 500.')
    bullet(doc, 'Do not make direct ServiceM8 job creation API calls from public form submissions — inbox email only.')

    # ── 5.10 BUILD & DEPLOY
    h2(doc, '5.10 — Build & Deploy Commands')

    h3(doc, 'Development')
    code_block(doc,
        "npm install           # Install dependencies (first time only)\n"
        "npm run dev           # Vite dev server at localhost:5173 (hot reload)\n"
        "npm run lint          # ESLint\n"
        "npm run format        # Prettier auto-format"
    )

    h3(doc, 'Production build')
    code_block(doc,
        "npm run build\n"
        "# Outputs:\n"
        "#   dist/rg-calculator.js   (minified React bundle)\n"
        "#   dist/rg-calculator.css  (styles)"
    )

    h3(doc, 'Copy assets to plugin folder (Windows PowerShell)')
    code_block(doc,
        "copy dist\\rg-calculator.js  wordpress-plugin\\rg-calculator\\assets\\\n"
        "copy dist\\rg-calculator.css wordpress-plugin\\rg-calculator\\assets\\\n"
        "# DO NOT copy dist\\*.jpg — images stay in assets permanently"
    )

    h3(doc, 'Package for WordPress upload')
    code_block(doc,
        "Compress-Archive -Path wordpress-plugin\\rg-calculator `\n"
        "  -DestinationPath wordpress-plugin\\rg-calculator.zip -Force"
    )

    h3(doc, 'Verify build output')
    code_block(doc,
        "# Check bundle size\n"
        "ls -la dist/rg-calculator.js dist/rg-calculator.css\n"
        "# Expected: JS ~200-400 KB, CSS ~2 KB"
    )

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 6 — TROUBLESHOOTING
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 6 — Common Issues & Troubleshooting')

    issues = [
        (
            'Calculator does not appear on the page',
            [
                'Verify the shortcode [rg_calculator] is on the page.',
                'View page source — look for <div id="rg-calculator-root">. If missing, the shortcode is not rendering.',
                'Check the plugin is active: WP Admin → Plugins.',
                'Check browser console (F12) for JavaScript errors.',
                'Check that rg-calculator.js and rg-calculator.css are present in wp-content/plugins/rg-calculator/assets/.',
            ]
        ),
        (
            'REST API returns 404 (form submission fails)',
            [
                'WordPress permalinks need flushing.',
                'Fix: WP Admin → Settings → Permalinks → click Save Changes (no changes needed).',
                'Verify by visiting /wp-json/royal-glass/v1/pricing directly in your browser — should return JSON.',
            ]
        ),
        (
            'Admin notification email not received',
            [
                'Check spam/junk folder first.',
                'Check RG_LEAD_NOTIFY_EMAIL is correctly set in wp-config.php.',
                'PHP mail() may be blocked on Bluehost — install WP Mail SMTP and configure Gmail/Sendgrid.',
                'Test with a plugin like "Check Email" to verify wp_mail() is working at all.',
                'Check PHP error logs (cPanel → Error Logs) for mail-related errors.',
            ]
        ),
        (
            'Customer email not received',
            [
                'Check spam/junk folder — branded HTML emails can trigger spam filters.',
                'Same SMTP setup as admin email above.',
                'Verify the customer email address was captured correctly in WP Admin → Leads.',
            ]
        ),
        (
            'ServiceM8 email not received in inbox',
            [
                'Check RG_SM8_INBOX_EMAIL is defined in wp-config.php.',
                'Check the SM8 inbox address is correct (find it in ServiceM8 → Settings → Inbox).',
                'Check PHP error logs for servicem8_status = failed in the leads table.',
                'Confirm general email delivery works (see admin email troubleshooting above).',
            ]
        ),
        (
            'Turnstile CAPTCHA blocks legitimate submissions',
            [
                'Verify both keys are correct in wp-config.php (site key and secret must match the same Cloudflare domain).',
                'Ensure the site domain matches the domain registered in Cloudflare Turnstile settings.',
                'Temporarily comment out both constants to disable Turnstile for testing — revert immediately after.',
                'Check browser console for Turnstile widget errors.',
            ]
        ),
        (
            'Pricing changes in admin do not show on the calculator',
            [
                'Hard refresh the calculator page (Ctrl+Shift+R) — browser may have cached the old response.',
                'Check GET /wp-json/royal-glass/v1/pricing returns the updated values.',
                'If a caching plugin (WP Rocket, W3TC) is active, flush its cache.',
            ]
        ),
        (
            'Database error on lead submission',
            [
                'Check WP Admin → Tools → Site Health for database issues.',
                'The wp_rg_leads table may not have been created. Deactivate and reactivate the plugin to re-run the activation hook.',
                'After a PHP update or plugin version jump, newer columns may be missing. The plugin auto-migrates columns on admin page load — visit WP Admin → RG Calculator → Leads once.',
            ]
        ),
        (
            'Address autocomplete does not suggest anything',
            [
                'The address lookup uses OpenStreetMap Nominatim — no API key required, but it is rate-limited.',
                'Try typing a suburb or street name (not just a number).',
                'Nominatim may be temporarily unavailable — the field still works for manual entry.',
                'Users can type the address manually if autocomplete fails.',
            ]
        ),
        (
            'Estimate shows 0 or no price',
            [
                'The pricing fetch from /pricing may have failed and fallen back to defaults.',
                'Check the browser network tab — look for a failed request to /wp-json/royal-glass/v1/pricing.',
                'If dev mode, pricing falls back to DEFAULT_PRICING in config.ts automatically.',
            ]
        ),
        (
            '"Too many requests" error on submission',
            [
                'Rate limit: 10 submissions per IP per hour.',
                'Wait an hour, or temporarily disable the rate limit in api.php for testing.',
                'Logged-in WordPress admins bypass the rate limit.',
            ]
        ),
        (
            'Plugin update breaks the calculator (JS errors)',
            [
                'Check that rg-calculator.js and rg-calculator.css were copied to the assets folder after the build.',
                'Compare file modification dates with the build output.',
                'Ensure you did not overwrite image assets (*.jpg) — they must remain in the assets folder.',
            ]
        ),
    ]

    for title, steps_list in issues:
        h3(doc, title)
        for s in steps_list:
            bullet(doc, s)

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 7 — KNOWN LIMITATIONS & FUTURE WORK
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 7 — Known Limitations & Future Work')

    h2(doc, 'Current limitations')
    two_col_table(doc, [
        ('No automated tests',
         'There is no test suite. Validation is manual. A developer adding features should test all wizard paths manually.'),
        ('Estimate not server-recalculated',
         'The estimate submitted with the lead is client-supplied (bounds-clamped but not recomputed server-side). '
         'A malicious user could submit an inflated-but-valid estimate. Low practical risk.'),
        ('/estimate-email missing full security',
         'The share estimate endpoint only has nonce + rate limit. Missing Turnstile, honeypot, and time gate. '
         'Low practical risk on Bluehost.'),
        ('Rate limit non-atomic',
         'The rate limit check/set uses separate read + write operations. Concurrent requests can exceed the limit by 1. '
         'Not exploitable in practice at this traffic level.'),
        ('marketingConsent not stored',
         'The marketing consent checkbox is captured in React state but never sent to the server or saved to the DB.'),
        ('No lead export',
         'Leads can only be viewed in the WP Admin UI — there is no CSV export. Developers can query wp_rg_leads directly.'),
        ('No pagination on leads list',
         'The admin leads page loads up to 200 leads. As lead volume grows, pagination should be added.'),
        ('Google Maps key unused',
         'RG_GOOGLE_MAPS_KEY is passed to the frontend but the address autocomplete now uses Nominatim. '
         'The constant can be removed when the relevant wp_localize_script line is cleaned up.'),
    ], header=['Limitation', 'Details'], col_widths=[2.0, 4.0])

    doc.add_paragraph()
    h2(doc, 'Planned / suggested future work')
    bullet(doc, 'Server-side estimate recalculation in api.php to eliminate client-trust dependency.')
    bullet(doc, 'Full Turnstile + honeypot protection on /estimate-email endpoint.')
    bullet(doc, 'CSV export for leads (admin-leads.php).')
    bullet(doc, 'Lead list pagination.')
    bullet(doc, 'Store marketingConsent to DB and respect it for future email.')
    bullet(doc, 'Automated unit tests for the pricing engine (engine.ts).')
    bullet(doc, 'Remove unused Google Maps constant and wp_localize_script field.')

    page_break(doc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PART 8 — QUICK REFERENCE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(doc, 'PART 8 — Quick Reference')

    h2(doc, 'Key file locations')
    two_col_table(doc, [
        ('Change prices (code)',         'src/lib/calculator/config.ts → DEFAULT_PRICING'),
        ('Change prices (live, no rebuild)', 'WP Admin → RG Calculator → Pricing'),
        ('Add/change wizard step options',   'src/components/CalculatorForm.tsx'),
        ('Change estimate email template',   'wordpress-plugin/rg-calculator/includes/email.php'),
        ('Add REST API endpoint',            'wordpress-plugin/rg-calculator/includes/api.php'),
        ('Change DB schema',                 'wordpress-plugin/rg-calculator/includes/database.php'),
        ('Add/change validation rules',      'wordpress-plugin/rg-calculator/includes/validation.php'),
        ('Plugin constants & shortcode',     'wordpress-plugin/rg-calculator/rg-calculator.php'),
        ('Image assets',                     'wordpress-plugin/rg-calculator/assets/*.jpg'),
    ], header=['Task', 'File'], col_widths=[2.8, 3.2])

    doc.add_paragraph()
    h2(doc, 'wp-config.php constants cheat sheet')
    code_block(doc,
        "define('RG_TURNSTILE_SITE_KEY', '...');  // Cloudflare public key\n"
        "define('RG_TURNSTILE_SECRET',   '...');  // Cloudflare private key\n"
        "define('RG_LEAD_NOTIFY_EMAIL',  '...');  // Admin notification recipient\n"
        "define('RG_SM8_INBOX_EMAIL',    '...');  // ServiceM8 inbox (or comment to disable)"
    )

    h2(doc, 'REST API quick reference')
    two_col_table(doc, [
        ('GET  /pricing',          'Returns live PricingConfig JSON — no auth required'),
        ('POST /leads',            'Submit a new lead — full security pipeline'),
        ('POST /estimate-email',   'Forward estimate to email — nonce + rate limit'),
    ], header=['Endpoint', 'Purpose'], col_widths=[2.5, 3.5])

    doc.add_paragraph()
    h2(doc, 'Deploy checklist')
    numbered(doc, 'npm run build')
    numbered(doc, 'Copy dist/rg-calculator.js and dist/rg-calculator.css to assets/ (not images)')
    numbered(doc, 'Package ZIP (Compress-Archive)')
    numbered(doc, 'Upload to WordPress → Plugins → Upload')
    numbered(doc, 'Activate plugin')
    numbered(doc, 'Verify /estimate page loads the calculator')
    numbered(doc, 'Submit a test lead and confirm all three emails arrive')
    numbered(doc, 'Check WP Admin → Leads shows the test submission')

    h2(doc, 'Contacts & resources')
    two_col_table(doc, [
        ('Royal Glass website',         'royalglass.co.nz'),
        ('Calculator estimate page',    'royalglass.co.nz/estimate'),
        ('WordPress admin',             'royalglass.co.nz/wp-admin'),
        ('Cloudflare Turnstile',        'dash.cloudflare.com → Turnstile'),
        ('ServiceM8',                   'app.servicem8.com'),
        ('Plugin git repository',       'Private — ask the development team'),
    ], header=['Resource', 'URL / Location'], col_widths=[2.5, 3.5])

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # FOOTER
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.add_paragraph()
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        f'Royal Glass Cost Calculator Plugin v2.3.0  ·  '
        f'Generated {datetime.date.today().strftime("%d %B %Y")}  ·  Confidential'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    run.font.italic = True

    # ── Save
    out_path = 'RG_Calculator_Documentation.docx'
    doc.save(out_path)
    print(f'Saved: {out_path}')


if __name__ == '__main__':
    build_doc()
